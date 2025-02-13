import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import unicodedata

st.set_page_config(page_title="PLU Listen Anwendung")

def normalize_string(s):
    return ''.join(c for c in unicodedata.normalize('NFKD', s) if not unicodedata.combining(c))

def generate_plu_list(mother_file, plu_week_file):
    """
    Erstellt eine PLU-Liste mit Kategorien auf neuen Seiten.
    """
    mother_file = pd.ExcelFile(mother_file)
    plu_week_df = pd.read_excel(plu_week_file, header=None, names=['PLU'])
    
    # Entferne ungültige Werte und wandle in Integer um
    plu_week_df = plu_week_df.dropna(subset=["PLU"])
    plu_week_df["PLU"] = pd.to_numeric(plu_week_df["PLU"], errors='coerce').dropna().astype(int)
    
    categories = mother_file.sheet_names
    filtered_data = []
    
    for category in categories:
        category_data = mother_file.parse(category)
        if "PLU" not in category_data.columns or "Artikel" not in category_data.columns:
            continue  # Überspringe Kategorien, die nicht die benötigten Spalten enthalten
        
        matched_data = pd.merge(plu_week_df, category_data, on="PLU", how="inner")
        if matched_data.empty:
            continue  # Überspringe leere Datensätze
        
        matched_data["Artikel_normalized"] = matched_data["Artikel"].apply(normalize_string)
        matched_data = matched_data.sort_values(by="Artikel_normalized").reset_index(drop=True)
        matched_data.drop(columns=["Artikel_normalized"], inplace=True)
        matched_data["Kategorie"] = category
        filtered_data.append(matched_data)
    
    if not filtered_data:
        raise ValueError("Keine passenden Daten gefunden.")
    
    combined_df = pd.concat(filtered_data, ignore_index=True).drop_duplicates(subset=['PLU'])
    doc = Document()
    
    first_page = True
    for category, data in combined_df.groupby("Kategorie"):
        if not first_page:
            doc.add_page_break()
        first_page = False
        doc.add_heading(category, level=1)
        
        for _, row in data.iterrows():
            doc.add_paragraph(f"{row['PLU']}\t{row['Artikel']}")
    
    output_word = BytesIO()
    doc.save(output_word)
    output_word.seek(0)
    
    return output_word

st.title("PLU-Listen Generator")

uploaded_mother_file = st.file_uploader("Optionale Mutterdatei hochladen (Excel)", type="xlsx")
uploaded_plu_week_file = st.file_uploader("PLU-Wochen-Datei hochladen (Excel)", type="xlsx")

if st.button("PLU-Liste generieren"):
    if uploaded_plu_week_file:
        try:
            mother_file = uploaded_mother_file if uploaded_mother_file else "mother_file.xlsx"
            
            with st.spinner("Processing..."):
                output_word = generate_plu_list(mother_file, uploaded_plu_week_file)
            
            st.success("PLU-Liste erfolgreich erstellt!")
            st.download_button(
                label="PLU-Liste herunterladen (Word)",
                data=output_word,
                file_name="plu_list.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except FileNotFoundError:
            st.error("Die Standard-Mutterdatei wurde nicht gefunden. Bitte laden Sie eine eigene Mutterdatei hoch.")
        except ValueError as e:
            st.error(f"Eingabefehler: {str(e)}")
        except Exception as e:
            st.error(f"Ein unerwarteter Fehler ist aufgetreten: {str(e)}")
    else:
        st.warning("Bitte laden Sie die PLU-Wochen-Datei hoch.")

# Neuer Datenschutzhinweis
st.markdown("⚠️ **Hinweis:** Diese Anwendung speichert keine Daten und hat keinen Zugriff auf Ihre Dateien.")
st.markdown("🌟 **Erstellt von Christoph R. Kaiser mit Hilfe von Künstlicher Intelligenz.**")
