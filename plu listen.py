import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO

def generate_plu_list(mother_file, plu_week_file):
    """
    Diese Funktion erstellt eine PLU-Liste basierend auf den hochgeladenen Mutterdateien und PLU-Woche-Dateien.

    Parameter:
    - mother_file: Hochgeladene Excel-Datei (Mutterdatei) als BytesIO.
    - plu_week_file: Hochgeladene Excel-Datei (PLU-Woche) als BytesIO.

    Rückgabe:
    - BytesIO-Objekt mit der generierten Word-Datei.
    """
    # 1. Daten laden
    mother_file = pd.ExcelFile(mother_file)
    plu_week_df = pd.read_excel(plu_week_file)

    # Sicherstellen, dass PLU-Nummern Integer sind
    plu_week_df["PLU"] = plu_week_df["PLU"].astype(int)

    # Kategorien aus der Mutterdatei laden
    categories = mother_file.sheet_names
    filtered_data = {}

    # 2. Abgleich der PLU-Nummern und Filtern der Artikel
    for category in categories:
        category_data = mother_file.parse(category)
        matched_data = pd.merge(plu_week_df, category_data, on="PLU", how="inner")
        matched_data = matched_data.sort_values(by="Artikel").reset_index(drop=True)
        filtered_data[category] = matched_data

    # 3. Word-Dokument erstellen
    doc = Document()

    for category, data in filtered_data.items():
        # Kategorie als Überschrift
        doc.add_heading(category, level=1)

        # PLU-Nummern und Artikel einfügen
        for _, row in data.iterrows():
            doc.add_paragraph(f"{row['PLU']}\t{row['Artikel']}")

    # 4. Dokument in BytesIO speichern
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# Streamlit App
st.title("PLU List Generator")

# Datei-Uploads
uploaded_mother_file = st.file_uploader("Upload Mother File (Excel)", type="xlsx")
uploaded_plu_week_file = st.file_uploader("Upload PLU Week File (Excel)", type="xlsx")

if st.button("Generate PLU List"):
    if uploaded_mother_file and uploaded_plu_week_file:
        try:
            with st.spinner("Processing..."):
                # PLU-Liste generieren
                output_file = generate_plu_list(uploaded_mother_file, uploaded_plu_week_file)

            # Download-Link für die Datei anzeigen
            st.success("PLU List successfully generated!")
            st.download_button(
                label="Download PLU List",
                data=output_file,
                file_name="plu_list.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception as e:
            st.error(f"An error occurred: {str(e)}")
    else:
        st.warning("Please upload both the Mother File and the PLU Week File.")
